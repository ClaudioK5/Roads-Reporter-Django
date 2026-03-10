from django.shortcuts import render
from .forms import ExcelUploadForm
import pandas as pd
from docx import Document
import os
from django.conf import settings


def upload_excel(request):
    message = None
    error = None
    total_length = None
    total_text = None
    report_url = None

    expected_columns = [
        "№ п/п",
        "Наименование",
        "Значение автомобильной дороги",
        "Категория",
        "Протяженность, км"
    ]

    if request.method == 'POST':
        form = ExcelUploadForm(request.POST, request.FILES)

        if form.is_valid():
            uploaded_file = request.FILES['excel_file']

            try:
                df = pd.read_excel(uploaded_file)

                # Нормализуем названия колонок
                df.columns = [
                    " ".join(str(col).replace("\n", " ").split())
                    for col in df.columns
                ]

                if list(df.columns) != expected_columns:
                    error = (
                        "Неверная структура Excel-файла. "
                        "Проверьте названия и порядок колонок."
                    )
                else:
                    df = df[expected_columns]

                    # Обрабатываем русские десятичные числа
                    df["Протяженность, км"] = (
                        df["Протяженность, км"]
                        .astype(str)
                        .str.replace(",", ".", regex=False)
                        .str.strip()
                    )

                    df["Протяженность, км"] = pd.to_numeric(
                        df["Протяженность, км"],
                        errors="coerce"
                    )

                    total_length = round(df["Протяженность, км"].sum(), 2)
                    total_text = f"Общая протяженность автомобильных дорог составляет {total_length} км"

                    # ---------- ГЕНЕРАЦИЯ WORD ----------
                    document = Document()

                    document.add_paragraph(
                        "Таблица 1 - Перечень и характеристика автомобильных дорог, "
                        "проходящих по территории муниципального округа"
                    )

                    table = document.add_table(rows=1, cols=len(expected_columns))
                    table.style = "Table Grid"

                    # Заголовки таблицы
                    hdr_cells = table.rows[0].cells
                    for i, col_name in enumerate(expected_columns):
                        hdr_cells[i].text = col_name

                    # Данные таблицы
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, value in enumerate(row):
                            row_cells[i].text = str(value)

                    document.add_paragraph("")
                    document.add_paragraph(total_text)

                    # Создаём папку reports, если её нет
                    reports_dir = os.path.join(settings.MEDIA_ROOT, "reports")
                    os.makedirs(reports_dir, exist_ok=True)

                    report_filename = "road_report.docx"
                    report_path = os.path.join(reports_dir, report_filename)

                    document.save(report_path)

                    report_url = settings.MEDIA_URL + "reports/" + report_filename

                    message = f'Файл "{uploaded_file.name}" успешно загружен и обработан.'

            except Exception as e:
                error = f"Ошибка при обработке файла: {str(e)}"
    else:
        form = ExcelUploadForm()

    return render(request, 'roads/upload.html', {
        'form': form,
        'message': message,
        'error': error,
        'total_length': total_length,
        'total_text': total_text,
        'report_url': report_url
    })